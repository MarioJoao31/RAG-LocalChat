import os
import fitz  # PyMuPDF
import docx

def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_pdf(file_path):
    doc = fitz.open(file_path)
    return "\n".join([page.get_text() for page in doc])

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_documents(directory="Data/Documents"):
    """
    Read the folder and loads all text, pdf, and docx files.
    Returns a list of dictionaries with 'text' and 'path' keys.
    """
    docs = []
    for root, _, files in os.walk(directory):
        for file in files:
            path = os.path.join(root, file)
            ext = file.lower().split('.')[-1]
            try:
                if ext == 'txt':
                    text = read_txt(path)
                elif ext == 'pdf':
                    text = read_pdf(path)
                elif ext == 'docx':
                    text = read_docx(path)
                else:
                    continue
                docs.append({'text': text, 'path': path})
            except Exception as e:
                print(f"Failed to read {file}: {e}")
    return docs

def save_uploaded_file(file, directory="Docs/Documentation"):
    """
    Saves an uploaded file (from Streamlit, for example) into the specified directory.
    Returns the full path where the file was saved.
    """
    if not os.path.exists(directory):
        os.makedirs(directory)

    file_path = os.path.join(directory, file.name)
    with open(file_path, "wb") as f:
        f.write(file.getbuffer())
    return file_path

def read_file(file):
    """
    Reads the content of a file-like object (.txt, .pdf, .docx).
    Accepts file-like objects (e.g., from Streamlit's uploader).
    Returns the extracted text.
    """
    ext = file.name.split(".")[-1].lower()

    if ext == "txt":
        return file.read().decode("utf-8")

    elif ext == "pdf":
        file_bytes = file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join([page.get_text() for page in doc])

    elif ext == "docx":
        docx_io = io.BytesIO(file.read())
        docx_file = docx.Document(docx_io)
        return "\n".join([p.text for p in docx_file.paragraphs])

    else:
        raise ValueError(f"Unsupported file format: {ext}")